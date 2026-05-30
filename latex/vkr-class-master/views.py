import os
import json
import re
import tempfile
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PyPDF2
import fitz

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django import forms
from books.models import Book, Arc, Chapter, ChapterVersion
import requests
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_http_methods
from django.conf import settings


# Загрузка PDF/DOCX файла и получение информации о страницах
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def upload_file(request):
    try:
        uploaded_file = request.FILES.get('file')
        chapter_id = request.POST.get('chapter_id')
        if not uploaded_file:
            return JsonResponse({'error': 'Файл не найден'}, status=400)
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
            for chunk in uploaded_file.chunks():
                tmp_file.write(chunk)
            tmp_path = tmp_file.name
        request.session['import_file_path'] = tmp_path
        request.session['import_file_name'] = uploaded_file.name
        total_pages = 0
        preview = ""
        if uploaded_file.name.endswith('.pdf'):
            with open(tmp_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                total_pages = len(pdf_reader.pages)
                if total_pages > 0:
                    preview = pdf_reader.pages[0].extract_text()[:500]
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(tmp_path)
            full_text = ''
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text += p.text.strip() + ' '
            total_chars = len(full_text)
            chars_per_page = 2000
            total_pages = max(1, (total_chars + chars_per_page - 1) // chars_per_page)
            preview = full_text[:500]
        return JsonResponse({
            'success': True,
            'total_pages': total_pages,
            'preview': preview,
            'file_name': uploaded_file.name
        })
    except Exception as e:
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


# Импорт выбранных страниц из PDF/DOCX в главу
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def import_pages(request):
    try:
        data = json.loads(request.body)
        chapter_id = data.get('chapter_id')
        pages = data.get('pages', [])
        chapter = get_object_or_404(Chapter, id=chapter_id, arc__book__user=request.user)
        file_path = request.session.get('import_file_path')
        file_name = request.session.get('import_file_name', '')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'Файл не найден. Загрузите файл заново.'}, status=400)
        imported_text = []
        if file_name.endswith('.pdf'):
            import fitz
            doc = fitz.open(file_path)
            for page_num in sorted(pages):
                if 1 <= page_num <= len(doc):
                    page = doc[page_num - 1]
                    text = page.get_text("text")
                    if text and text.strip():
                        text = text.replace('\n', ' ')
                        text = re.sub(r' {2,}', ' ', text).strip()
                        if text:
                            imported_text.append(f'<p>{text}</p>')
            doc.close()
        elif file_name.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            full_text = ''
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text += p.text.strip() + ' '
            total_chars = len(full_text)
            chars_per_page = 2000
            for page_num in sorted(pages):
                start_char = (page_num - 1) * chars_per_page
                end_char = min(start_char + chars_per_page, total_chars)
                if start_char >= total_chars:
                    return JsonResponse({
                        'error': f'Страница {page_num} выходит за пределы документа (всего символов: {total_chars})'
                    }, status=400)
                page_text = full_text[start_char:end_char]
                if page_text.strip():
                    sentences = re.split(r'(?<=[.!?])\s+', page_text)
                    html_parts = []
                    current_para = []
                    for sentence in sentences:
                        current_para.append(sentence)
                        if len(current_para) >= 3:
                            html_parts.append(f'<p>{" ".join(current_para)}</p>')
                            current_para = []
                    if current_para:
                        html_parts.append(f'<p>{" ".join(current_para)}</p>')
                    imported_text.append('\n'.join(html_parts))
        content = '\n'.join(imported_text) if imported_text else '<p>Нет текста для импорта</p>'
        return JsonResponse({
            'success': True,
            'content': content,
            'pages_count': len(pages)
        })
    except Exception as e:
        print(f"Import error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


# Вызов AI ассистента через OpenRouter
@csrf_exempt
@require_http_methods(["POST"])
def ai_chat(request):
    try:
        data = json.loads(request.body)
        user_message = data.get('message', '')
        context = data.get('context', {})
        is_book_mode = context.get('isBookMode', True)
        if is_book_mode:
            book_title = context.get('bookTitle', 'Неизвестная книга')
            chapter_title = context.get('chapterTitle', 'Неизвестная глава')
            current_content = context.get('currentChapterContent', '')
            all_chapters = context.get('allChaptersContent', '')
            clean_current = re.sub(r'<[^>]+>', '', current_content)[:3000] if current_content else ''
            clean_all = re.sub(r'<[^>]+>', '', all_chapters)[:3000] if all_chapters else ''
            system_prompt = f"""Ты — профессиональный AI ассистент для писателя.

Контекст:
- Книга: "{book_title}"
- Текущая глава: "{chapter_title}"
- Содержание главы: {clean_current[:2000]}
- Контекст всей книги: {clean_all[:2000]}"""
        else:
            system_prompt = """Ты — полезный и точный AI ассистент. Строго следуй правилам:

ПРАВИЛА ОТВЕТОВ:
1. Отвечай ТОЛЬКО если уверен в информации
2. Если не знаешь точного ответа — честно скажи: "Я не знаю точного ответа"
3. Не выдумывай факты, цифры, даты
4. Для творческих вопросов — предлагай варианты, но помечай их как идеи
5. Для технических вопросов — давай точные инструкции
6. Отвечай на русском языке
7. Используй Markdown для структурирования ответа"""
        FREE_MODELS = [
            {"name": "tencent/hy3-preview", "limit": 100, "description": "Tencent Hy3 (бесплатно до 8 мая)"},
            {"name": "minimax/minimax-m2.5", "limit": 100, "description": "MiniMax M2.5"},
            {"name": "google/gemma-4-31b-it:free", "limit": 100, "description": "Google Gemma 4 31B"},
            {"name": "openai/gpt-oss-120b:free", "limit": 100, "description": "OpenAI GPT-OSS 120B"},
            {"name": "nvidia/nemotron-3-super:free", "limit": 100, "description": "NVIDIA Nemotron 3 Super"},
            {"name": "openrouter/free", "limit": 50, "description": "Автовыбор (запасной вариант)"}
        ]
        headers = {
            "Authorization": f"Bearer {settings.OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": request.build_absolute_uri('/'),
            "X-Title": "Writer AI Assistant"
        }
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message}
        ]
        last_error = None
        used_model = None
        for model_config in FREE_MODELS:
            model_name = model_config["name"]
            payload = {
                "model": model_name,
                "messages": messages,
                "temperature": 0.3 if not is_book_mode else 0.7,
                "max_tokens": 4000,
                "top_p": 0.9,
                "frequency_penalty": 0.3,
                "presence_penalty": 0.1
            }
            try:
                response = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=60
                )
                if response.status_code == 200:
                    result = response.json()
                    reply = result['choices'][0]['message']['content']
                    used_model = model_name
                    return JsonResponse({'reply': reply, 'model': model_name})
                elif response.status_code == 429:
                    last_error = f"Лимит исчерпан для {model_name}, пробую следующую..."
                    print(f" {last_error}")
                    continue
                elif response.status_code == 402:
                    last_error = f"Модель {model_name} требует оплату, пропускаем..."
                    print(f" {last_error}")
                    continue
                else:
                    error_data = response.json() if response.text else {}
                    error_msg = error_data.get('error', {}).get('message', 'Неизвестная ошибка')
                    last_error = f"Ошибка {model_name}: {error_msg}"
                    print(f" {last_error}")
                    continue
            except requests.Timeout:
                last_error = f"Таймаут для {model_name}, пробую следующую..."
                print(f" {last_error}")
                continue
            except Exception as e:
                last_error = f"Ошибка {model_name}: {str(e)}"
                print(f" {last_error}")
                continue
        return JsonResponse({
            'reply': f""" Все бесплатные модели временно недоступны.

**Последняя ошибка**: {last_error}

**Что делать**:
- Подождите несколько часов (лимиты сбрасываются)
- Проверьте API ключ в настройках
- Попробуйте позже (бесплатные модели имеют ограничения)"""
        }, status=500)
    except Exception as e:
        print(f"AI ошибка: {e}")
        return JsonResponse({'reply': f" Техническая ошибка: {str(e)}"}, status=500)

# Получение содержимого всех глав книги для контекста AI
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def get_book_content(request):
    try:
        data = json.loads(request.body)
        book_id = data.get('book_id')
        book = get_object_or_404(Book, id=book_id, user=request.user)
        all_content = []
        for arc in book.arcs.all().order_by('arc_number'):
            all_content.append(f"\n=== {arc.title_arc} ===\n")
            for chapter in arc.chapters.all().order_by('chapter_number'):
                clean_text = re.sub(r'<[^>]+>', '', chapter.content or '')
                all_content.append(f"## {chapter.title_chapter}")
                all_content.append(clean_text[:500])
                all_content.append("---")
        content_text = "\n".join(all_content)
        return JsonResponse({'success': True, 'content': content_text})
    except Book.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Книга не найдена'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


# Вход пользователя в систему
def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            return redirect('dashboard')
    return render(request, 'main/auth.html')


from django.contrib import messages

# Регистрация нового пользователя
def register_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')
        confirm = request.POST.get('confirm_password')
        if password != confirm:
            messages.error(request, 'Пароли не совпадают')
            return redirect('register')
        if User.objects.filter(username=username).exists():
            messages.error(request, 'Пользователь с таким именем уже существует')
            return redirect('register')
        if User.objects.filter(email=email).exists():
            messages.error(request, 'Пользователь с таким email уже существует')
            return redirect('register')
        user = User.objects.create_user(username=username, email=email, password=password)
        login(request, user)
        return redirect('dashboard')
    return render(request, 'main/register.html')


# Выход из системы
def logout_view(request):
    logout(request)
    return redirect('login')


# Главная страница со списком книг
@login_required
def dashboard_view(request):
    books = Book.objects.filter(user=request.user)
    for book in books:
        if book.cover_path:
            if book.cover_path.name == 'covers/default.png':
                continue
            full_path = os.path.join(settings.MEDIA_ROOT, book.cover_path.name)
            if not os.path.isfile(full_path):
                book.cover_path = 'covers/default.png'
                book.save()
        else:
            book.cover_path = 'covers/default.png'
            book.save()
    return render(request, 'main/dashboard.html', {'books': books})


# Обновление обложки книги
@login_required
def update_cover(request, book_id):
    book = get_object_or_404(Book, id=book_id, user=request.user)
    if request.method == 'POST' and request.FILES.get('cover'):
        cover_file = request.FILES['cover']
        if not cover_file.content_type.startswith('image/'):
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'error': 'Можно загружать только изображения'}, status=400)
            return redirect('book_info', book_id=book.id)
        if book.cover_path and book.cover_path.name != 'covers/default.png':
            old_cover_path = os.path.join(settings.MEDIA_ROOT, book.cover_path.name)
            if os.path.exists(old_cover_path):
                os.remove(old_cover_path)
        book.cover_path = cover_file
        book.save()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'cover_url': book.cover_path.url})
        return redirect('book_info', book_id=book.id)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'error': 'Файл не выбран'}, status=400)
    return redirect('book_info', book_id=book.id)


# Удаление обложки книги
@login_required
def delete_cover(request, book_id):
    book = get_object_or_404(Book, id=book_id, user=request.user)
    if book.cover_path and book.cover_path.name != 'covers/default.png':
        full_path = os.path.join(settings.MEDIA_ROOT, book.cover_path.name)
        if os.path.exists(full_path):
            os.remove(full_path)
    book.cover_path = 'covers/default.png'
    book.save()
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'success': True, 'cover_url': book.cover_path.url})
    return redirect('book_info', book_id=book.id)

# Страница информации о книге
@login_required
def book_info(request, book_id):
    book = get_object_or_404(Book, id=book_id, user=request.user)
    return render(request, 'main/book_info.html', {'book': book})

# Форма редактирования книги
class BookForm(forms.ModelForm):
    class Meta:
        model = Book
        fields = ['title_book', 'size', 'status']


# Форма добавления/редактирования книги
class AddBookForm(forms.ModelForm):
    SIZE_CHOICES = [
        ('мини', 'Мини'),
        ('миди', 'Миди'),
        ('макси', 'Макси'),
    ]
    STATUS_CHOICES = [
        ('черновик', 'Черновик'),
        ('заброшено', 'Заброшено'),
        ('завершено', 'Завершено'),
        ('заморожено', 'Заморожено'),
    ]
    size = forms.ChoiceField(choices=SIZE_CHOICES)
    status = forms.ChoiceField(choices=STATUS_CHOICES)
    title_book = forms.CharField(label='Название книги', max_length=200)

    class Meta:
        model = Book
        fields = ['title_book', 'size', 'status']


# Добавление или редактирование книги
@login_required
def add_book(request):
    book_id = request.GET.get('id')
    book = None
    if book_id:
        book = get_object_or_404(Book, id=book_id, user=request.user)
    if request.method == 'POST':
        if book:
            form = AddBookForm(request.POST, instance=book)
        else:
            form = AddBookForm(request.POST)
        if form.is_valid():
            book = form.save(commit=False)
            book.user = request.user
            if not book_id:
                book.page_count = 0
                book.chapter_count = 0
            book.save()
            return redirect('book_info', book_id=book.id)
    else:
        if book:
            form = AddBookForm(instance=book)
        else:
            form = AddBookForm()
    return render(request, 'main/add_book.html', {'form': form, 'book': book})


# Добавление новой главы в арку
@login_required
@require_http_methods(["POST"])
def add_chapter(request, arc_id):
    try:
        data = json.loads(request.body or "{}")
        chapter_title = data.get('title')
        if not chapter_title:
            return JsonResponse({'error': 'Название главы обязательно'}, status=400)
        arc = get_object_or_404(Arc, id=arc_id, book__user=request.user)
        last_chapter = arc.chapters.order_by('-chapter_number').first()
        if last_chapter:
            chapter_number = last_chapter.chapter_number + 1
        else:
            chapter_number = 1
        chapter = Chapter.objects.create(
            title_chapter=chapter_title,
            arc=arc,
            content='',
            chapter_number=chapter_number
        )
        return JsonResponse({'id': chapter.id, 'title': chapter.title_chapter})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный JSON'}, status=400)
    except Exception as e:
        print("ERROR:", e)
        return JsonResponse({'error': str(e)}, status=500)

# Добавление новой арки в книгу
@login_required
@require_http_methods(["POST"])
def add_arc(request, book_id):
    try:
        data = json.loads(request.body or "{}")
        arc_title = data.get('title')
        if not arc_title:
            return JsonResponse({'error': 'Название арки обязательно'}, status=400)
        book = get_object_or_404(Book, id=book_id, user=request.user)
        last_arc = book.arcs.order_by('-arc_number').first()
        if last_arc:
            arc_number = last_arc.arc_number + 1
        else:
            arc_number = 1
        arc = Arc.objects.create(
            title_arc=arc_title,
            book=book,
            arc_number=arc_number
        )
        return JsonResponse({'id': arc.id, 'title': arc.title_arc})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный JSON'}, status=400)
    except Exception as e:
        print("ERROR:", e)
        return JsonResponse({'error': str(e)}, status=500)

# Удаление арки
@require_http_methods(["DELETE"])
def delete_arc(request, arc_id):
    try:
        arc = Arc.objects.get(id=arc_id)
        arc.delete()
        return JsonResponse({'success': True})
    except Arc.DoesNotExist:
        return JsonResponse({'error': 'Арка не найдена'}, status=404)

# Удаление главы
@csrf_exempt
@require_http_methods(["DELETE"])
def delete_chapter(request, chapter_id):
    try:
        chapter = Chapter.objects.get(id=chapter_id)
        arc_id = chapter.arc.id
        chapter.delete()
        return JsonResponse({'success': True, 'arc_id': arc_id})
    except Chapter.DoesNotExist:
        return JsonResponse({'error': 'Глава не найдена'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# Страница чтения/редактирования главы
def chapter_info(request, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id)
    chapters_in_arc = list(chapter.arc.chapters.all().order_by('chapter_number', 'id'))
    current_index = -1
    for i, ch in enumerate(chapters_in_arc):
        if ch.id == chapter.id:
            current_index = i
            break
    previous_chapter = chapters_in_arc[current_index - 1] if current_index > 0 else None
    next_chapter = chapters_in_arc[current_index + 1] if current_index < len(chapters_in_arc) - 1 else None
    context = {
        'chapter': chapter,
        'previous_chapter': previous_chapter,
        'next_chapter': next_chapter,
    }
    return render(request, 'main/chapter.html', context)


# Сохранение содержимого главы с версионированием
@login_required
@require_http_methods(["POST"])
def save_chapter(request, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id, arc__book__user=request.user)
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "invalid json"}, status=400)
    content = data.get("content", "")
    if chapter.content != content:
        ChapterVersion.objects.create(chapter=chapter, content=chapter.content)
        chapter.content = content
        chapter.save()
    return JsonResponse({"success": True})


# Извлечение RGB цвета из style строки
def parse_rgb(style_string):
    if not style_string:
        return None
    match = re.search(r'color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', style_string)
    if match:
        return RGBColor(int(match[1]), int(match[2]), int(match[3]))
    match = re.search(r'color:\s*#([A-Fa-f0-9]{6})', style_string)
    if match:
        hex_color = match[1]
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
    return None


# Очистка Quill HTML от классов редактора
def clean_quill_html(html_content):
    if not html_content:
        return ""
    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.find_all(class_=True):
        classes = element.get("class", [])
        current_style = element.get("style", "")
        if "ql-align-center" in classes:
            current_style += "text-align: center; "
        if "ql-align-right" in classes:
            current_style += "text-align: right; "
        if "ql-align-justify" in classes:
            current_style += "text-align: justify; "
        if "ql-align-left" in classes:
            current_style += "text-align: left; "
        if current_style.strip():
            element["style"] = current_style.strip()
        new_classes = [c for c in classes if not c.startswith("ql-")]
        if new_classes:
            element["class"] = new_classes
        elif "class" in element.attrs:
            del element.attrs["class"]
    for i in range(1, 7):
        for tag in soup.find_all(f"h{i}"):
            tag.name = f"h{i}"
            if not tag.get("style"):
                sizes = {1: 24, 2: 20, 3: 18, 4: 16, 5: 14, 6: 12}
                tag["style"] = f"font-size: {sizes[i]}pt; font-weight: bold;"
    return str(soup)


# Применение форматирования списков в DOCX
def apply_list_formatting(paragraph, list_type, level=0):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    if list_type == 'ordered':
        numId.set(qn('w:val'), '1')
    else:
        numId.set(qn('w:val'), '2')
    numPr.append(numId)
    pPr.append(numPr)
    indent = OxmlElement('w:ind')
    indent.set(qn('w:left'), str(level * 720 + 360))
    pPr.append(indent)


# Конвертация HTML содержимого в документ Word
def add_html_content_to_doc(doc, html_content):
    if not html_content:
        return
    clean_html = clean_quill_html(html_content)
    soup = BeautifulSoup(clean_html, "html.parser")
    for el in soup.find_all(["p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol"]):
        text_content = el.get_text(strip=True)
        if not text_content and not el.find_all(["img", "br"]):
            continue
        if el.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(el.name[1])
            p = doc.add_paragraph()
            align_style = el.get("style", "")
            if "text-align: center" in align_style:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "text-align: right" in align_style:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            text = el.get_text().strip()
            run = p.add_run(text)
            run.bold = True
            font_sizes = {1: 24, 2: 20, 3: 18, 4: 16, 5: 14, 6: 12}
            run.font.size = Pt(font_sizes[level])
            continue
        if el.name in ["ul", "ol"]:
            list_type = 'ordered' if el.name == 'ol' else 'unordered'
            for li in el.find_all("li", recursive=False):
                li_text = li.get_text().strip()
                if not li_text:
                    continue
                p = doc.add_paragraph()
                level = 0
                parent_li = li.find_parent("li")
                if parent_li:
                    level = 1
                apply_list_formatting(p, list_type, level)
                def process_li_child(child, paragraph):
                    for subchild in child.children:
                        if subchild.name is None and subchild.string:
                            run = paragraph.add_run(subchild.string.strip())
                        elif subchild.name in ["strong", "b"]:
                            run = paragraph.add_run(subchild.get_text())
                            run.bold = True
                        elif subchild.name in ["em", "i"]:
                            run = paragraph.add_run(subchild.get_text())
                            run.italic = True
                        elif subchild.name == "u":
                            run = paragraph.add_run(subchild.get_text())
                            run.underline = True
                        elif subchild.name == "span":
                            text = subchild.get_text()
                            if text:
                                run = paragraph.add_run(text)
                                color = parse_rgb(subchild.get("style", ""))
                                if color:
                                    run.font.color.rgb = color
                        elif subchild.name == "ul" or subchild.name == "ol":
                            for nested_li in subchild.find_all("li"):
                                nested_p = doc.add_paragraph()
                                nested_list_type = 'ordered' if subchild.name == 'ol' else 'unordered'
                                apply_list_formatting(nested_p, nested_list_type, level + 1)
                                for nested_child in nested_li.children:
                                    if nested_child.name is None and nested_child.string:
                                        nested_p.add_run(nested_child.string.strip())
                                    elif nested_child.name in ["strong", "b"]:
                                        run = nested_p.add_run(nested_child.get_text())
                                        run.bold = True
                                    elif nested_child.name in ["em", "i"]:
                                        run = nested_p.add_run(nested_child.get_text())
                                        run.italic = True
                process_li_child(li, p)
            continue
        p = doc.add_paragraph()
        align_style = el.get("style", "")
        if "text-align: center" in align_style:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "text-align: right" in align_style:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "text-align: justify" in align_style:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        def process_element(element, paragraph):
            for child in element.children:
                if child.name is None:
                    if child.string and child.string.strip():
                        run = paragraph.add_run(child.string.strip())
                    continue
                if child.name == "br":
                    paragraph.add_run("\n")
                    continue
                if child.name == "img":
                    continue
                if child.string and child.string.strip():
                    text = child.string.strip()
                    run = paragraph.add_run(text)
                    if child.name in ["strong", "b"]:
                        run.bold = True
                    if child.name in ["em", "i"]:
                        run.italic = True
                    if child.name == "u":
                        run.underline = True
                    child_style = child.get("style", "")
                    color = parse_rgb(child_style)
                    if color:
                        run.font.color.rgb = color
                if child.find_all(recursive=False):
                    process_element(child, paragraph)
        process_element(el, p)

# Экспорт одной главы в DOCX
def export_docx(request, chapter_id):
    chapter = Chapter.objects.get(id=chapter_id)
    doc = Document()
    add_html_content_to_doc(doc, chapter.content)
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{chapter.title_chapter}.docx"'
    doc.save(response)
    return response

# Экспорт выбранных глав в один DOCX файл
def export_selected_chapters(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            chapter_ids = data.get('chapter_ids', [])
            if not chapter_ids:
                return JsonResponse({'error': 'Не выбрано ни одной главы'}, status=400)
            chapters = Chapter.objects.filter(id__in=chapter_ids).order_by('arc__id', 'id')
            if not chapters.exists():
                return JsonResponse({'error': 'Главы не найдены'}, status=404)
            doc = Document()
            book_title = chapters.first().arc.book.title_book
            title = doc.add_heading(f'{book_title}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for chapter in chapters:
                doc.add_heading(chapter.title_chapter, level=1)
                add_html_content_to_doc(doc, chapter.content)
                doc.add_page_break()
            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response["Content-Disposition"] = f'attachment; filename="{book_title}_chapters.docx"'
            doc.save(response)
            return response
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Метод не разрешен'}, status=405)


# Переименование арки
@require_http_methods(["POST"])
def rename_arc(request, arc_id):
    try:
        data = json.loads(request.body)
        new_title = data.get('title', '').strip()
        if not new_title:
            return JsonResponse({'error': 'Название арки не может быть пустым'}, status=400)
        arc = Arc.objects.get(id=arc_id)
        arc.title_arc = new_title
        arc.save()
        return JsonResponse({'success': True, 'id': arc.id, 'title': arc.title_arc})
    except Arc.DoesNotExist:
        return JsonResponse({'error': 'Арка не найдена'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# Переименование главы
@require_http_methods(["POST"])
def rename_chapter(request, chapter_id):
    try:
        data = json.loads(request.body)
        new_title = data.get('title')
        if not new_title:
            return JsonResponse({'error': 'Название не может быть пустым'}, status=400)
        chapter = Chapter.objects.get(id=chapter_id)
        chapter.title_chapter = new_title
        chapter.save()
        return JsonResponse({'title': chapter.title_chapter})
    except Chapter.DoesNotExist:
        return JsonResponse({'error': 'Глава не найдена'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# Удаление книги
@login_required
@require_http_methods(["DELETE"])
def delete_book(request, book_id):
    try:
        book = Book.objects.get(id=book_id, user=request.user)
        book.delete()
        return JsonResponse({'success': True})
    except Book.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Книга не найдена'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)